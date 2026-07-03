# -*- coding: utf-8 -*-
"""
Generates synthetic input data for the demo project "Automatische Dokumentenerstellung":
  1. input/Teilnehmerliste.xlsx  - participant list (Name, Kurs, Datum, Stunden)
  2. input/Vorlage_Teilnahmebescheinigung.docx - Word template with {{PLACEHOLDER}} fields

All names and companies are fictional. Run:  python generate_input.py
"""

import os
import random
from datetime import date

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

random.seed(42)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_DIR = os.path.join(BASE_DIR, "input")
os.makedirs(INPUT_DIR, exist_ok=True)

# ---------------------------------------------------------------- participants
VORNAMEN = ["Lena", "Jonas", "Katharina", "Felix", "Miriam", "Tobias", "Annika",
            "Sebastian", "Claudia", "Martin", "Sophie", "Daniel", "Franziska", "Patrick"]
NACHNAMEN = ["Brandt", "Keller", "Vogel", "Neumann", "Krüger", "Seidel", "Böhm",
             "Winkler", "Lorenz", "Haas", "Engel", "Pohl", "Wenzel", "Franke"]
KURSE = [
    ("Excel Grundlagen für den Büroalltag", 16),
    ("Datenschutz am Arbeitsplatz (DSGVO)", 8),
    ("Erste Hilfe im Betrieb", 9),
    ("Professionelle E-Mail-Kommunikation", 6),
    ("Arbeitssicherheit und Brandschutz", 8),
]

teilnehmer = []
namen_pool = random.sample([f"{v} {n}" for v in VORNAMEN for n in NACHNAMEN], 12)
for name in namen_pool:
    kurs, stunden = random.choice(KURSE)
    monat = random.randint(3, 6)
    tag = random.randint(1, 28)
    teilnehmer.append((name, kurs, date(2026, monat, tag), stunden))
teilnehmer.sort(key=lambda t: t[2])

# ---------------------------------------------------------------- Excel list
wb = Workbook()
ws = wb.active
ws.title = "Teilnehmer"

header_fill = PatternFill("solid", fgColor="1F4E79")
header_font = Font(bold=True, color="FFFFFF")
thin = Side(style="thin", color="BFBFBF")
border = Border(left=thin, right=thin, top=thin, bottom=thin)

headers = ["Name", "Kurs", "Datum", "Stunden"]
for col, h in enumerate(headers, 1):
    c = ws.cell(row=1, column=col, value=h)
    c.fill = header_fill
    c.font = header_font
    c.border = border
    c.alignment = Alignment(horizontal="center")

for r, (name, kurs, datum, stunden) in enumerate(teilnehmer, 2):
    ws.cell(row=r, column=1, value=name).border = border
    ws.cell(row=r, column=2, value=kurs).border = border
    dc = ws.cell(row=r, column=3, value=datum)
    dc.number_format = "DD.MM.YYYY"
    dc.border = border
    sc = ws.cell(row=r, column=4, value=stunden)
    sc.border = border
    sc.alignment = Alignment(horizontal="center")

ws.column_dimensions["A"].width = 24
ws.column_dimensions["B"].width = 42
ws.column_dimensions["C"].width = 14
ws.column_dimensions["D"].width = 10

xlsx_path = os.path.join(INPUT_DIR, "Teilnehmerliste.xlsx")
wb.save(xlsx_path)
print(f"OK: {xlsx_path} ({len(teilnehmer)} Zeilen)")

# ---------------------------------------------------------------- Word template
doc = Document()
for section in doc.sections:
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(12)

DUNKELBLAU = RGBColor(0x1F, 0x4E, 0x79)
GRAU = RGBColor(0x7F, 0x7F, 0x7F)

# fictional training company as letterhead
kopf = doc.add_paragraph()
kopf.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = kopf.add_run("Schulungszentrum Elbtal GmbH")
r.font.size = Pt(14)
r.font.bold = True
r.font.color.rgb = DUNKELBLAU

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Musterstraße 12 · 01067 Dresden · info@sz-elbtal.example")
r.font.size = Pt(9)
r.font.color.rgb = GRAU

doc.add_paragraph()
doc.add_paragraph()

titel = doc.add_paragraph()
titel.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = titel.add_run("TEILNAHMEBESCHEINIGUNG")
r.font.size = Pt(22)
r.font.bold = True
r.font.color.rgb = DUNKELBLAU

doc.add_paragraph()

hiermit = doc.add_paragraph()
hiermit.alignment = WD_ALIGN_PARAGRAPH.CENTER
hiermit.add_run("Hiermit wird bestätigt, dass")

name_p = doc.add_paragraph()
name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = name_p.add_run("{{NAME}}")
r.font.size = Pt(18)
r.font.bold = True

doc.add_paragraph()

text_p = doc.add_paragraph()
text_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
text_p.add_run("am {{DATUM}} erfolgreich an der Schulung")

kurs_p = doc.add_paragraph()
kurs_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = kurs_p.add_run("„{{KURS}}“")
r.font.size = Pt(14)
r.font.bold = True

umfang_p = doc.add_paragraph()
umfang_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
umfang_p.add_run("im Umfang von {{STUNDEN}} Unterrichtsstunden teilgenommen hat.")

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

ort_p = doc.add_paragraph()
ort_p.add_run("Dresden, den {{AUSSTELLUNGSDATUM}}")

doc.add_paragraph()
doc.add_paragraph()

unterschrift = doc.add_paragraph()
unterschrift.add_run("____________________________")
rolle = doc.add_paragraph()
r = rolle.add_run("Schulungsleitung")
r.font.size = Pt(10)
r.font.color.rgb = GRAU

docx_path = os.path.join(INPUT_DIR, "Vorlage_Teilnahmebescheinigung.docx")
doc.save(docx_path)
print(f"OK: {docx_path}")
